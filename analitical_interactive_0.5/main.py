"""
Secondary interface:
- data postprocessing;
- generating main GUI;
- generating report (in development);
"""
import tkinter
from tkinter import ttk
from main_data_processing import *
from tkinter import *
from docx.enum.text import WD_BREAK
import os

sko_result_intens = {}
sko_result_pos = {}

pars_zone_numbers_local = 20
pars_zone_len_local = 25

sko_pos = []
sko_int = []
abs_pos = []

differences = {}
diff = []

something = main_window_destroy
reload_x = 0
secondary_window = None
secondary_window_destroy = False

selected_error = ''
error_for_visualisation = ''
first_init = True
indexes = []
first_index_red = False


def message_for_frames():
    """Calculates minimum, average and maximum error values for output to the main frame"""

    global sko_int
    global sko_pos
    global abs_pos
    global something
    global sko_result_intens
    global sko_result_pos

    # first - minimum, second - average, third - maximum error for each variable
    sko_pos = [[0, 1000], 0.0, [0, 0]]
    sko_int = [[0, 1000], 0.0, [0, 0]]
    abs_pos = [[0, 1000], 0.0, [0, 0]]

    # min\max\avr in standard deviation of intensities
    list_for_min_max_avr = []
    for pos, data in sko_result_intens.items():
        if data != 'None information' and sko_result_intens[pos]['metrology_var'] == 1:
            list_for_min_max_avr.append(sko_result_intens[pos]['sko'])
        else:
            list_for_min_max_avr.append(1964)
        compare = min(list_for_min_max_avr)
        compare_2 = max(list_for_min_max_avr)
        if compare < sko_int[0][1]:
            sko_int[0].clear()
            sko_int[0].append(pos)
            sko_int[0].append(compare)
        if compare_2 > sko_int[2][1] and compare_2 != 1964:
            sko_int[2].clear()
            sko_int[2].append(pos)
            sko_int[2].append(compare_2)
        while 1964 in list_for_min_max_avr:
            list_for_min_max_avr.remove(1964)
        sko_int[1] = 0
        for things in list_for_min_max_avr:
            sko_int[1] += things
        if len(list_for_min_max_avr) > 0:
            sko_int[1] = sko_int[1] / (len(list_for_min_max_avr))

    # min\max\avr in standard deviation of positions
    list_for_min_max_avr = []
    for pos, data in sko_result_pos.items():
        if data != 'None information' and sko_result_pos[pos]['metrology_var'] == 1 and data['exp_data']:
            list_for_min_max_avr.append(sko_result_pos[pos]['sqroot'])
        else:
            list_for_min_max_avr.append(1964)
        compare = min(list_for_min_max_avr)
        compare_2 = max(list_for_min_max_avr)
        if compare < sko_pos[0][1]:
            sko_pos[0].clear()
            sko_pos[0].append(pos)
            sko_pos[0].append(compare)
        if compare_2 > sko_pos[2][1] and compare_2 != 1964:
            sko_pos[2].clear()
            sko_pos[2].append(pos)
            sko_pos[2].append(compare_2)
        while 1964 in list_for_min_max_avr:
            list_for_min_max_avr.remove(1964)
        sko_pos[1] = 0
        for things in list_for_min_max_avr:
            sko_pos[1] += things
        if len(list_for_min_max_avr) > 0:
            sko_pos[1] = sko_pos[1] / (len(list_for_min_max_avr))

    # min\max\avr in absolute error
    list_for_min_max_avr = []
    for pos, data in absolute_error.items():
        if 'None information' in data or not data['exp_data'] or data['metrology_var'] == 0:
            list_for_min_max_avr.append(0.000001)
        else:
            list_for_min_max_avr.append(max(data['exp_data']))
        abs_pos[2] = [peak_hkl[list_for_min_max_avr.index(max(list_for_min_max_avr))], max(list_for_min_max_avr)]

    list_for_min_max_avr = []
    for pos, data in absolute_error.items():
        if 'None information' not in data and data['metrology_var'] == 1 and pos != str(peak_hkl[0]):
            list_for_min_max_avr.append(min(data['exp_data']))
        else:
            list_for_min_max_avr.append(1964)
        abs_pos[0] = [peak_hkl[list_for_min_max_avr.index(min(list_for_min_max_avr))], min(list_for_min_max_avr)]
    summ = 0
    len_timer = 0
    for pos, data in absolute_error.items():
        if 'None information' not in data and 0.0 not in data and data['exp_data'] and data['metrology_var'] == 1:
            for items in data['exp_data']:
                summ += items
                len_timer += 1
    if len_timer == 0:
        abs_pos[1] = 0.0
    elif len(list_for_min_max_avr) > 0:
        abs_pos[1] = summ / len_timer

    # if only one element for processing left
    if abs_pos[0][1] == 1964:
        abs_pos[0][1] = abs_pos[2][1]

    something = True


def change_search_area():
    """Creates two sliders to select the number of rows and columns to be read from Excel file; writes it in global"""

    def write_pars_numbers(self):
        global pars_zone_numbers_local
        pars_zone_numbers_local = (pars_zone.get())
        print(self)

    def write_pars_len(self):
        global pars_zone_len_local
        pars_zone_len_local = (pars_len.get())
        print(self)

    search_area_button.destroy()
    pars_zone = IntVar()
    scale_numbers = Scale(frame1, from_=10, to=200, orient=HORIZONTAL, label='Количество строк: ',
                          resolution=10, length=140, variable=pars_zone,
                          command=write_pars_numbers)
    scale_numbers.set(20)
    scale_numbers.pack()

    pars_len = IntVar()
    scale_len = Scale(frame1, from_=25, to=250, orient=HORIZONTAL, label='Количество столбцов: ',
                      resolution=25, length=140, variable=pars_len,
                      command=write_pars_len)
    scale_len.set(50)
    scale_len.pack()


def open_file():
    """Executes the program body (calculation of deviations and errors)"""

    global sko_result_pos
    global sko_result_intens
    global reload_x
    global secondary_window
    global secondary_window_destroy
    global error_for_visualisation
    error_for_visualisation = ''

    if main_window_destroy:
        pars_location(pars_zone_len_local)
        trigger = file_location(directory_expand)
        if trigger:
            return None
    pars(pars_zone_numbers_local)

    if check_for_no_answer:
        main_window.destroy()
        return None

    # basic calculations and conversions
    sko_result_pos = sko_init(error_calculated_positions, peak_hkl)
    sko_result_intens = sko_init(error_calculated_intensities, peak_hkl)
    abs_err()
    message_for_frames()

    # removes the initial windows after selecting a file
    if main_window_destroy and not secondary_window_destroy:
        main_window.destroy()
    if secondary_window_destroy:
        secondary_window.destroy()

    secondary_window = Tk()
    secondary_window.title("Metrology")

    # disables the ability to zoom the page
    secondary_window.resizable(False, False)

    main_frame = LabelFrame(secondary_window)
    main_frame.pack(side=TOP)

    # creates the basic frames
    frame_for_checkbutton = LabelFrame(main_frame, bg='#dddddd')
    frame_for_down_buttons = LabelFrame(main_frame)

    # for standard deviation of intensities
    frame3 = LabelFrame(main_frame)
    # for standard deviation of positions
    frame4 = LabelFrame(main_frame)
    # for absolute error
    frame5 = LabelFrame(main_frame)

    frame_for_checkbutton.pack(side=RIGHT)
    frame_for_down_buttons.pack(side=BOTTOM)

    frame3.pack(side=TOP)
    frame4.pack(side=TOP)
    frame5.pack(side=TOP)

    # frames for standard deviation of intensities
    frame3_1 = LabelFrame(frame3)
    frame3_2 = LabelFrame(frame3)
    frame3_3 = LabelFrame(frame3)
    frame3_4 = LabelFrame(frame3)
    frame3_1.pack(side=TOP)
    frame3_2.pack(side=LEFT)
    frame3_3.pack(side=LEFT)
    frame3_4.pack(side=LEFT)

    # frames for standard deviation of positions
    frame4_1 = LabelFrame(frame4)
    frame4_2 = LabelFrame(frame4)
    frame4_3 = LabelFrame(frame4)
    frame4_4 = LabelFrame(frame4)
    frame4_1.pack(side=TOP)
    frame4_2.pack(side=LEFT)
    frame4_3.pack(side=LEFT)
    frame4_4.pack(side=LEFT)

    # frames for absolute error
    frame5_1 = LabelFrame(frame5)
    frame5_2 = LabelFrame(frame5)
    frame5_3 = LabelFrame(frame5)
    frame5_4 = LabelFrame(frame5)
    frame5_1.pack(side=TOP)
    frame5_2.pack(side=LEFT)
    frame5_3.pack(side=LEFT)
    frame5_4.pack(side=LEFT)

    hkl_for_frame = []

    def update_errors_for_frames():
        """Updates data in metrology results by adding variable triggers for checkboxes"""

        # adds a check for outputs over values; a variable for checkboxes, and the color of the checkbox
        for keys, values in sko_result_intens.items():
            if values != 'None information' and values['sko'] > 2:
                values['checkbutton_variable'] = False
                values['checkbutton_variable_for_def'] = IntVar()
            elif values == 'None information':
                pass
            else:
                values['checkbutton_variable'] = True
                values['checkbutton_variable_for_def'] = IntVar()

        for keys, values in sko_result_pos.items():
            if values != 'None information' and round(values['sko']) >= 0.02 and values['metrology_var'] == 1:
                values['checkbutton_variable'] = False
                values['checkbutton_variable_for_def'] = IntVar()
            elif values == 'None information':
                pass
            else:
                values['checkbutton_variable'] = True
                values['checkbutton_variable_for_def'] = IntVar()

        timer_trig = 0
        for keys, values in absolute_error.items():
            if 'None information' not in values and values != [] and max(values['exp_data']) > 0.02:
                absolute_error[peak_hkl[timer_trig]]['checkbutton_variable'] = False
                absolute_error[peak_hkl[timer_trig]]['checkbutton_variable_for_def'] = IntVar()

            elif 'None information' in values or values == []:
                absolute_error[peak_hkl[timer_trig]] = 'None information'
            else:
                absolute_error[peak_hkl[timer_trig]]['checkbutton_variable'] = True
                absolute_error[peak_hkl[timer_trig]]['checkbutton_variable_for_def'] = IntVar()

            timer_trig += 1

    update_errors_for_frames()

    class Labels:
        """Frame configuration 'St.dev of intensities', 'St.dev of positions' and 'Absolute error'"""

        def __init__(self):
            # frame for standard deviation of intensities
            lbl_main_1 = Label(frame3_1, text='СКО Интенсивности:', width=63, bg='lightgreen')
            lbl_main_1.pack(side=TOP)

            lbl_min_3_1 = Label(frame3_2, text='минимальное:', width=20, bg='grey60')
            lbl_min_3_1.pack(side=TOP)
            self.lbl_real_min_3_1 = Label(frame3_2, width=20, bg='white')
            self.lbl_real_min_3_1.pack(side=TOP)
            self.lbl_hkl_min_3_1 = Label(frame3_2, width=20, bg='grey70')
            self.lbl_hkl_min_3_1.pack(side=TOP)

            lbl_avr_3_1 = Label(frame3_3, text='среднее:', width=20, bg='grey60')
            lbl_avr_3_1.pack(side=TOP)
            self.lbl_real_avr_3_1 = Label(frame3_3, width=20, bg='white')
            self.lbl_real_avr_3_1.pack(side=TOP)
            lbl_hkl_avr_3_1 = Label(frame3_3, text='<-- hkl -->', width=20, bg='grey70')
            lbl_hkl_avr_3_1.pack(side=TOP)

            lbl_max_3_1 = Label(frame3_4, text='максимальное:', width=20, bg='grey60')
            lbl_max_3_1.pack(side=TOP)
            self.lbl_real_max_3_1 = Label(frame3_4, width=20)
            self.lbl_real_max_3_1.pack(side=TOP)
            self.lbl_hkl_max_3_1 = Label(frame3_4, width=20, bg='grey70')
            self.lbl_hkl_max_3_1.pack(side=TOP)

            # frame for standard deviation of positions
            lbl_main_2 = Label(frame4_1, text='СКО Положения:', width=63, bg='lightgreen')
            lbl_main_2.pack(side=TOP)

            lbl_min_4_1 = Label(frame4_2, text='минимальное:', width=20, bg='grey60')
            lbl_min_4_1.pack(side=TOP)
            self.lbl_real_min_4_1 = Label(frame4_2, width=20, bg='white')
            self.lbl_real_min_4_1.pack(side=TOP)
            self.lbl_hkl_min_4_1 = Label(frame4_2, width=20, bg='grey70')
            self.lbl_hkl_min_4_1.pack(side=TOP)

            lbl_avr_4_1 = Label(frame4_3, text='среднее:', width=20, bg='grey60')
            lbl_avr_4_1.pack(side=TOP)
            self.lbl_real_avr_4_1 = Label(frame4_3, width=20, bg='white')
            self.lbl_real_avr_4_1.pack(side=TOP)
            lbl_hkl_avr_4_1 = Label(frame4_3, text='<-- hkl -->', width=20, bg='grey70')
            lbl_hkl_avr_4_1.pack(side=TOP)

            lbl_max_4_1 = Label(frame4_4, text='максимальное:', width=20, bg='grey60')
            lbl_max_4_1.pack(side=TOP)
            self.lbl_real_max_4_1 = Label(frame4_4, width=20)
            self.lbl_real_max_4_1.pack(side=TOP)
            self.lbl_hkl_max_4_1 = Label(frame4_4, width=20, bg='grey70')
            self.lbl_hkl_max_4_1.pack(side=TOP)

            # frame for absolute error
            lbl_main_3 = Label(frame5_1, text='Абсолютная погрешность:', width=63, bg='lightgreen')
            lbl_main_3.pack(side=TOP)

            lbl_min_5_1 = Label(frame5_2, text='минимальное:', width=20, bg='grey60')
            lbl_min_5_1.pack(side=TOP)
            self.lbl_real_min_5_1 = Label(frame5_2, width=20, bg='white')
            self.lbl_real_min_5_1.pack(side=TOP)
            self.lbl_hkl_min_5_1 = Label(frame5_2, width=20, bg='grey70')
            self.lbl_hkl_min_5_1.pack(side=TOP)

            lbl_avr_5_1 = Label(frame5_3, text='среднее:', width=20, bg='grey60')
            lbl_avr_5_1.pack(side=TOP)
            self.lbl_real_avr_5_1 = Label(frame5_3, width=20, bg='white')
            self.lbl_real_avr_5_1.pack(side=TOP)
            lbl_hkl_avr_5_1 = Label(frame5_3, text='<-- hkl -->', width=20, bg='grey70')
            lbl_hkl_avr_5_1.pack(side=TOP)

            lbl_max_5_1 = Label(frame5_4, text='максимальное:', width=20, bg='grey60')
            lbl_max_5_1.pack(side=TOP)
            self.lbl_real_max_5_1 = Label(frame5_4, width=20)
            self.lbl_real_max_5_1.pack(side=TOP)
            self.lbl_hkl_max_5_1 = Label(frame5_4, width=20, bg='grey70')
            self.lbl_hkl_max_5_1.pack(side=TOP)

        def labels_config(self):
            """Frame generation 'St.dev of intensities', 'St.dev of positions' and 'Absolute error'"""

            # variables to change the color and font of the maximum frame values when going out of bounds
            sko_int_color = 'green2'
            sko_pos_color = 'green2'
            sko_avr_color = 'green2'
            sko_int_color_txt = 'black'
            sko_pos_color_txt = 'black'
            sko_avr_color_txt = 'black'

            if sko_pos[2][1] == 0 and sko_int[2][1] == 0:
                abs_pos[0][1] = 0.0
                abs_pos[1] = 0.0
                abs_pos[2][1] = 0.0
                abs_pos[0][0] = 0.0
                abs_pos[2][0] = 0.0
                sko_pos[0][1] = 0.0
                sko_pos[1] = 0.0
                sko_pos[2][1] = 0.0
                sko_int[0][1] = 0.0
                sko_int[1] = 0.0
                sko_int[2][1] = 0.0

            else:
                if sko_int[2][1] > 2:
                    sko_int_color = 'red'
                    sko_int_color_txt = 'white'
                if sko_pos[2][1] > 0.02:
                    sko_pos_color = 'red'
                    sko_pos_color_txt = 'white'
                if abs_pos[2][1] > 0.02:
                    sko_avr_color = 'red'
                    sko_avr_color_txt = 'white'

            self.lbl_real_min_3_1.config(text=str(round(sko_int[0][1], 3)) + '%')
            self.lbl_hkl_min_3_1.config(text=str(sko_int[0][0]))
            self.lbl_real_avr_3_1.config(text=str(round(float(sko_int[1]), 3)) + '%')
            self.lbl_real_max_3_1.config(text=str(round(sko_int[2][1], 3)) + '%', bg=sko_int_color,
                                         fg=sko_int_color_txt)
            self.lbl_hkl_max_3_1.config(text=str(sko_int[2][0]))

            self.lbl_real_min_4_1.config(text=str(round(sko_pos[0][1], 4)))
            self.lbl_hkl_min_4_1.config(text=str(sko_pos[0][0]))
            self.lbl_real_avr_4_1.config(text=str(round(float(sko_pos[1]), 4)))
            self.lbl_real_max_4_1.config(text=str(round(sko_pos[2][1], 4)), bg=sko_pos_color, fg=sko_pos_color_txt)
            self.lbl_hkl_max_4_1.config(text=str(sko_pos[2][0]))

            self.lbl_real_min_5_1.config(text=str(round(abs_pos[0][1], 5)))
            self.lbl_hkl_min_5_1.config(text=str(abs_pos[0][0]))
            self.lbl_real_avr_5_1.config(text=str(round(float(abs_pos[1]), 5)))
            self.lbl_real_max_5_1.config(text=str(round(abs_pos[2][1], 5)), bg=sko_avr_color, fg=sko_avr_color_txt)
            self.lbl_hkl_max_5_1.config(text=str(abs_pos[2][0]))

    reload_x = Labels()
    reload_x.labels_config()

    def reload():
        """
        Initiates the process of recalculation of the values and
        updating the information in the graphical interface
        """
        global first_init
        global error_for_visualisation
        if first_init:
            error_for_visualisation = 'all'
            first_init = False

        def reload_errors():
            """Recalculates the absolute error, with new values of checkbuttons"""

            # update boolean variable (state) values by clicking checkbuttons
            for keys, values in absolute_error.items():
                if values != 'None information':
                    # receives information one by one about all states of the buttons
                    absolute_error[keys]['metrology_var'] = absolute_error[keys][
                        'checkbutton_variable_for_def'].get()
                    sko_result_pos[keys]['metrology_var'] = absolute_error[keys][
                        'checkbutton_variable_for_def'].get()
                    sko_result_intens[keys]['metrology_var'] = absolute_error[keys][
                        'checkbutton_variable_for_def'].get()

            # finding all active keys
            absolute_error_checked = []
            for keys, values in absolute_error.items():
                if values != 'None information' and values is not None and values['metrology_var'] == 1:
                    absolute_error_checked.append(keys)

            # converting active keys into indexes to refer to the original data
            global indexes
            indexes = []
            for hkls in absolute_error_checked:
                indexes.append(peak_hkl.index(hkls))

            # recalculates absolute error
            absolute_error_before_unification = {}
            from main_data_processing import data_position
            global diff
            for name, peak_list in data_position.items():
                diff = []

                # resets the indices of the first element from which absolute error counts

                for index in indexes:
                    if index == indexes[0]:
                        absolute_error[peak_hkl[index]]['metrology_var'] = 0
                    else:
                        absolute_error[peak_hkl[index]]['metrology_var'] = 1

                    difference_experimental = (float(peak_list[index]) - float(peak_list[indexes[0]]))
                    difference_nist = (float(nist_peak_position[index])
                                       - float(nist_peak_position[indexes[0]]))
                    difference = abs(difference_nist - difference_experimental)
                    diff.append(difference)
                absolute_error_before_unification[name] = diff

            # updates absolute error data
            for index in indexes:
                absolute_error[peak_hkl[index]]['exp_data'] = []
                for items in absolute_error_before_unification.values():
                    absolute_error[peak_hkl[index]]['exp_data'].append(items[indexes.index(index)])

            # makes checkbuttons red if the absolute error is exceeded
            if error_for_visualisation == 'absolute error':
                timer_trig = 0
                for keys, values in absolute_error.items():
                    if 'None information' not in values and values != [] and max(values['exp_data']) >= 0.02:
                        absolute_error[peak_hkl[timer_trig]]['checkbutton_variable'] = False
                        globals()['chckbtn%d' % timer_trig].configure(fg='red')
                    elif 'None information' in values or values == []:
                        absolute_error[peak_hkl[timer_trig]] = 'None information'
                    else:
                        absolute_error[peak_hkl[timer_trig]]['checkbutton_variable'] = True
                        globals()['chckbtn%d' % timer_trig].configure(fg='black')

                    # if the checkbutton is released, makes it black
                    if 'None information' not in values and values != [] and not values['metrology_var']:
                        globals()['chckbtn%d' % timer_trig].configure(fg='black')
                    timer_trig += 1

                # if there is an excess in absolute error, makes the first checkbutton red
                global first_index_red
                first_index_red = False
                for keys, values in absolute_error.items():
                    if 'None information' not in values and values != []:
                        if not values['checkbutton_variable'] and values['metrology_var']:
                            first_index_red = True
                            break
                if first_index_red:
                    globals()['chckbtn%d' % indexes[0]].configure(fg='red')

            elif error_for_visualisation == 'positions deviation':
                # updates standard deviation of position

                timer_trig = 0
                for keys, values in sko_result_intens.items():
                    if 'None information' not in values and values != []:
                        globals()['chckbtn%d' % timer_trig].configure(fg='black')
                    timer_trig += 1

                timer_trig = 0
                for keys, values in sko_result_pos.items():
                    if 'None information' not in values and values != []:
                        if float(values['sqroot']) >= 0.2:
                            globals()['chckbtn%d' % timer_trig].configure(fg='red')
                        if not values['metrology_var']:
                            globals()['chckbtn%d' % timer_trig].configure(fg='black')
                    timer_trig += 1

            elif error_for_visualisation == 'intensities deviation':
                # updates standard deviation of intensities

                timer_trig = 0
                for keys, values in sko_result_intens.items():
                    if 'None information' not in values and values != []:
                        globals()['chckbtn%d' % timer_trig].configure(fg='black')
                    timer_trig += 1

                timer_trig = 0
                for keys, values in sko_result_intens.items():
                    if 'None information' not in values and values != []:
                        if float(values['sko']) >= 2:
                            globals()['chckbtn%d' % timer_trig].configure(fg='red')
                        if not values['metrology_var']:
                            globals()['chckbtn%d' % timer_trig].configure(fg='black')
                    timer_trig += 1

            elif error_for_visualisation == 'all':
                # updates all errors

                # makes all checkbuttons red
                timer_trig = 0
                for keys, values in sko_result_intens.items():
                    if 'None information' not in values and values != []:
                        globals()['chckbtn%d' % timer_trig].configure(fg='black')
                    timer_trig += 1

                # updates absolute error
                first_index_red = False
                timer_trig = 0
                for keys, values in absolute_error.items():
                    if 'None information' not in values and values != []:
                        if max(values['exp_data']) >= 0.02 and values['metrology_var']:
                            globals()['chckbtn%d' % timer_trig].configure(fg='red')
                        if not values['checkbutton_variable'] and values['metrology_var']:
                            if max(values['exp_data']) >= 0.02:
                                first_index_red = True
                    timer_trig += 1

                # because of the shifts, have to paint everything black again
                timer_trig = 0
                for keys, values in absolute_error.items():
                    if 'None information' not in values and values != [] and not values['metrology_var']:
                        globals()['chckbtn%d' % timer_trig].configure(fg='black')
                timer_trig += 1

                # makes the first value red if there is an active deviation
                if indexes:
                    if first_index_red:
                        globals()['chckbtn%d' % indexes[0]].configure(fg='red')
                    else:
                        globals()['chckbtn%d' % indexes[0]].configure(fg='black')

                #  updates standard deviation of intensities
                timer_trig = 0
                for keys, values in sko_result_intens.items():
                    if 'None information' not in values and values != []:
                        if values['sko'] >= 2 and values['metrology_var']:
                            globals()['chckbtn%d' % timer_trig].configure(fg='red')
                    timer_trig += 1

                # updates standard deviation of positions
                timer_trig = 0
                for keys, values in sko_result_pos.items():
                    if 'None information' not in values and values != []:
                        if values['sqroot'] >= 0.2 and values['metrology_var']:
                            globals()['chckbtn%d' % timer_trig].configure(fg='red')
                    timer_trig += 1

            else:
                # in case of chosen "none"
                timer_trig = 0
                for keys, values in sko_result_intens.items():
                    if 'None information' not in values and values != []:
                        globals()['chckbtn%d' % timer_trig].configure(fg='black')
                    timer_trig += 1

        reload_errors()
        message_for_frames()
        reload_x.labels_config()

    # check positions and converts to a single form
    from main_data_processing import nist_peak_position

    timer = 0
    max_len_of_key = []

    for key, value in sko_result_pos.items():
        max_len_of_key.append(len(key))
    max_len_of_key = max(max_len_of_key)

    for key, value in sko_result_pos.items():
        if len(key) < max_len_of_key:
            hkl_for_frame.append(' ' + key + ' ' * (int(max_len_of_key) - len(key) + 2))
        else:
            hkl_for_frame.append(' ' + key + ' ')

        # if there is no information, then the frame is deactivated, if over 100, then a little less gaps
        if len(value) <= 15:

            if float(value['exp_data'][1]) >= 100:
                globals()['chckbtn%d' % timer] = Checkbutton(frame_for_checkbutton,
                                                             text=hkl_for_frame[-1] + ' -   ' + nist_peak_position[
                                                                 timer] + '°',
                                                             bg='#dddddd',
                                                             variable=absolute_error[peak_hkl[timer]][
                                                                 'checkbutton_variable_for_def'],
                                                             command=reload)
                globals()['chckbtn%d' % timer].grid(row=timer, sticky=W, pady=0)
                Checkbutton.select(globals()['chckbtn%d' % timer])
            else:
                globals()['chckbtn%d' % timer] = Checkbutton(frame_for_checkbutton,
                                                             text=hkl_for_frame[-1] + ' -   ' + nist_peak_position[
                                                                 timer] + '°  ',
                                                             bg='#dddddd',
                                                             command=reload,
                                                             variable=absolute_error[peak_hkl[timer]][
                                                                 'checkbutton_variable_for_def'])
                globals()['chckbtn%d' % timer].grid(row=timer, sticky=W, pady=0)
                Checkbutton.select(globals()['chckbtn%d' % timer])
        else:
            if float(nist_peak_position[timer]) >= 100:
                globals()['chckbtn%d' % timer] = Checkbutton(frame_for_checkbutton,
                                                             text=hkl_for_frame[-1] + ' -   ' + nist_peak_position[
                                                                 timer] + '°',
                                                             bg='grey70', state=DISABLED)
                globals()['chckbtn%d' % timer].grid(row=timer, sticky=W, pady=0)
            else:
                globals()['chckbtn%d' % timer] = Checkbutton(frame_for_checkbutton,
                                                             text=hkl_for_frame[-1] + ' -   ' + nist_peak_position[
                                                                 timer] + '°  ', bg='grey70',
                                                             state=DISABLED)
                globals()['chckbtn%d' % timer].grid(row=timer, sticky=W, pady=0)

        timer += 1

    reload()

    # меню выбора ошибки
    global selected_error
    selected_error = tkinter.StringVar()
    error_combobox = ttk.Combobox(frame_for_checkbutton, textvariable=selected_error, width=18)
    error_combobox['values'] = ('absolute error', 'positions deviation', 'intensities deviation', 'all', 'none')
    error_combobox['state'] = 'readonly'
    error_combobox.set('select deviation type')
    error_combobox.grid(row=timer, sticky=W, pady=4)

    def error_selected(event):
        """ handle the error changed event """
        global error_for_visualisation
        error_for_visualisation = selected_error.get()
        reload()

    error_combobox.bind('<<ComboboxSelected>>', error_selected)

    class Info:
        """
        Performs the initial generation and subsequent updating of information in the window generated by
        pressing the 'information' button
        """

        def __init__(self):
            """Generates basic static information fields"""

            self.open_close_trigger = 0
            from main_data_processing import directory_expand
            from main_data_processing import srm_name
            from main_data_processing import number_of_diffractograms
            global secondary_window
            self.information_frame = LabelFrame(secondary_window)

            self.main_info_lbl = Label(self.information_frame, width=64, justify=LEFT, wraplength=435, anchor=NW,
                                       relief=RIDGE, text=' Путь к файлу:\n  ' + str(directory_expand) +
                                                          '\n\n Стандарт образца:\n  NIST ' + str(srm_name) +
                                                          '\n\n Количество снятых дифрактограмм:\n  ' + str(
                                                           number_of_diffractograms) +
                                                          '\n\n Версия формул:\n  апрель 2022', pady=8)

            self.main_info_lbl.pack(side=LEFT, anchor=NW, fill=Y)
            self.info_for_checkbuttons = Label(self.information_frame, width=20, bg='#dddddd')
            self.info_for_checkbuttons.pack(side=LEFT, fill=Y)
            self.black_chckbtn = Label(self.info_for_checkbuttons, bg='#dddddd', justify=CENTER, text='ЧЕРНЫЙ:',
                                       fg='black', relief=GROOVE)
            self.black_chckbtn.pack(side=TOP, fill=X)
            self.black_chckbtn_info = Label(self.info_for_checkbuttons, bg='#dddddd', justify=CENTER,
                                            text='активное поле в \nрамках нормального\nотклонения',
                                            fg='black', relief=GROOVE, borderwidth=1, pady=3)
            self.black_chckbtn_info.pack(side=TOP, fill=X)
            self.gray_chckbtn = Label(self.info_for_checkbuttons, bg='#dddddd', justify=CENTER, text='СЕРЫЙ:',
                                      fg='gray30', relief=GROOVE)
            self.gray_chckbtn.pack(side=TOP, fill=X)
            self.gray_chckbtn_info = Label(self.info_for_checkbuttons, bg='#dddddd', justify=CENTER,
                                           text='отсутствует\nинформация о пике', relief=GROOVE, borderwidth=1, pady=3)
            self.gray_chckbtn_info.pack(side=TOP, fill=X)
            self.red_chckbtn_info = Label(self.info_for_checkbuttons, bg='#dddddd', justify=CENTER,
                                          text='КРАСНЫЙ:', fg='red', width=18, relief=GROOVE)
            self.red_chckbtn_info.pack(side=TOP, fill=X)

            self.red_chckbtn_info = Label(self.info_for_checkbuttons, bg='#dddddd', justify=CENTER, width=18,
                                          relief=GROOVE, text='выход за max\nотклонение', borderwidth=1, pady=3)
            self.red_chckbtn_info.pack(side=TOP, fill=Y)

        def reload(self):
            """is responsible for the appearance/disappearance of the information field"""

            self.open_close_trigger += 1
            if self.open_close_trigger == 1:
                self.information_frame.pack(side=TOP)

            else:
                self.open_close_trigger = 0
                self.information_frame.pack_forget()

    info_reload = Info()

    def chose_another_file():
        """Performs the function of processing a new file by pressing the 'select new file' button"""

        global first_init
        global secondary_window_destroy
        secondary_window_destroy = True
        first_init = True
        open_file()

    def make_report():
        """Function for 'make report' button. Making report and saving it into the 'word' file"""

        from main_data_processing import number_of_diffractograms
        from tkinter.filedialog import asksaveasfile
        import docx
        records_table_1 = []

        file = [('Word file', '*.docx')]
        defaultext = '*.docx'
        filename = asksaveasfile(filetypes=file, defaultextension=defaultext, initialfile='Метрология_дифрактометра_№_')
        # report = docx.Document(filename)

        document = docx.Document()  # docs.Document()!!!!!!!!!!!!!!!!!!!!!!!!!!

        par1 = document.add_heading('    8.4 Определение метрологических характеристик')
        par2 = document.add_paragraph('    Количество измеренных дифрактограмм: ' + str(number_of_diffractograms))
        par3 = document.add_heading('    8.4.1 Определение пределов допускаемой абсолютной погрешности при '
                                    'измерении угловых положений дифракционных максимумов')

        def create_table(tables, header, rows, style='Table Grid'):
            cols_number = len(header)

            table = tables.add_table(rows=1, cols=cols_number)
            table.style = style

            hdr_cells = table.rows[0].cells
            for column_number in range(cols_number):
                hdr_cells[column_number].text = header[column_number]

            for row in rows:
                row_cells = table.add_row().cells
                for column_number in range(cols_number):
                    row_cells[column_number].text = str(row[column_number])

            return table

        list_of_indexes = []
        standard_deviation = []
        calculated_deviations = []
        peak_positions_for_abs_error = []
        standard_difference = []

        from main_data_processing import data_position

        # генерирует n- списков пиков по m- количеству дифрактограмм
        for x in range(len(indexes)):
            standard_deviation_itter = []
            for b in range(number_of_diffractograms):
                standard_deviation_itter.append(data_position['spectrum_' + str(b + 1) + '_position'][indexes[x]])
            standard_deviation.append(standard_deviation_itter)

        for x in range(number_of_diffractograms):
            calculated_deviations.append([])

        for keys, values in absolute_error.items():
            if 'None information' not in values and values != []:
                if peak_hkl.index(keys) > indexes[0] and values['metrology_var']:
                    # первый столбец с индексами:
                    list_of_indexes.append(keys + ' - ' + peak_hkl[indexes[0]])

                    # второй столбец с эталонными разностями
                    difference = round(float(nist_peak_position[peak_hkl.index(keys)]) -
                                       float(nist_peak_position[indexes[0]]), 3)
                    if difference != 0.0:
                        standard_difference.append(difference)

                    # столбцы с расчетными разностями
                    peak_positions_for_abs_error = []
                    if difference > 0:
                        for x in range(number_of_diffractograms):
                            calculated_deviations[x].append(
                                round((values['exp_data'][x] - absolute_error[peak_hkl[indexes[0]]]['exp_data'][x]), 4))

                            peak_positions_for_abs_error_itter = []
                            # сортирует положения пиков, создавая массивы по количеству дифрактограмм
                            for i in range(len(indexes)):
                                peak_positions_for_abs_error_itter.append(standard_deviation[i][x])
                            peak_positions_for_abs_error.append(peak_positions_for_abs_error_itter)

        peak_difference_for_abs_error = []

        for iterations in range(number_of_diffractograms):
            peak_difference_for_abs_error.append([])

        # создает массив с разностью снятых значений положений пиков
        for iterations in range(len(indexes)):
            for iteration in range(number_of_diffractograms):
                result = round(float(peak_positions_for_abs_error[iteration][iterations]) - float(
                    peak_positions_for_abs_error[iteration][0]), 4)
                if result != 0.0:
                    peak_difference_for_abs_error[iteration].append(result)

        # хрен его знает, чё это, проверить!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
        for iterations in range(len(indexes) - 1):
            records_by_rows = [list_of_indexes[iterations], standard_difference[iterations]]
            for i in range(number_of_diffractograms):
                records_by_rows.append(peak_difference_for_abs_error[i][iterations])
                records_by_rows.append(calculated_deviations[i][iterations])
            records_table_1.append(records_by_rows)

        # собирает все данные в один массив для создания таблицы
        records_table_01 = []
        for values in records_table_1:
            counter = 2
            for counts in range((len(values) - 2) // 2):
                records_table_01.append([values[0], values[1], values[counter], values[counter + 1]])
                counter += 2

        # преобразует массив данных к выводному виду
        header_timer = 0
        for i in range(number_of_diffractograms):
            header_timer += 1
            records_table_01_for_output = []
            table_timer = 0
            for x in range(len(indexes) - 1):
                records_table_01_for_output.append(records_table_01[x * number_of_diffractograms + header_timer - 1])
                table_timer += 1

            # создает таблицы и хэдеры для таблиц
            par4 = document.add_paragraph('\n')
            run = par4.add_run('    Таблица ' + str(header_timer) + '. Данные дифрактограммы №' + str(header_timer))

            headers = ['Разница между максимумами соответствующих пиков',
                       'Эталонное значение разности',
                       'Вычисленное значение разности для измеренной дифрактограммы',
                       'Отклонение показаний дифрактометра']

            table1 = create_table(document, headers, records_table_01_for_output)

        par5 = document.add_paragraph()
        run = par5.add_run(
            '\n    Абсолютная погрешность дифрактометра при измерении угловых положений дифракционных '
            'максимумов составляет ' + str(round(abs_pos[2][1], 5)) + '°').font.bold = True
        p = document.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        par6 = document.add_heading(
            '    8.4.2 Определение среднеквадратичного отклонения (СКО) случайной составляющей погрешности при '
            'измерении угловых позиций Брэгговских отражений.')
        par7 = document.add_paragraph('\n')
        run = par7.add_run('    Таблица ' + str(header_timer + 1) + '. Среднеквадратичное отклонение положений')

        # создает таблицу для ско положений
        header_2 = ['hkl']
        for x in range(number_of_diffractograms):
            header_2.append(str(x + 1) + ' измер.')
        header_2.append('СКО положений')

        positions_for_report = []
        for i in indexes:
            positions = [peak_hkl[i]]
            for x in range(number_of_diffractograms):
                positions.append(round(float(peak_positions_for_abs_error[x][indexes.index(i)]), 3))
            positions.append(round(sko_result_pos[peak_hkl[i]]['sqroot'], 5))
            positions_for_report.append(positions)

        table1 = create_table(document, header_2, positions_for_report)

        par8 = document.add_paragraph()
        run = par8.add_run(
            '\n    Значение СКО погрешности измерения угловых позиций Брэгговских отражений составляет ' + str(
                round(sko_pos[2][1], 5)) + '.').font.bold = True
        p = document.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        par9 = document.add_heading(
            '    8.4.3 Среднеквадратичное отклонение случайной составляющей (СКО) погрешности определения '
            'относительных интенсивностей')
        par10 = document.add_paragraph('\n')
        run = par10.add_run('    Таблица ' + str(header_timer + 2) + '. Среднеквадратичное отклонение интенсивностей')
        header_3 = ['hkl']

        for x in range(number_of_diffractograms):
            header_3.append(str(x + 1) + ' измер.')
        header_3.append('СКО интенсивностей')

        intensities_for_report = []
        for i in indexes:
            intensities = [peak_hkl[i]]
            for x in range(number_of_diffractograms):
                intensities.append(sko_result_intens[peak_hkl[i]]['exp_data'][x])
            intensities.append(round(sko_result_intens[peak_hkl[i]]['sko'], 4))
            intensities_for_report.append(intensities)

        table1 = create_table(document, header_3, intensities_for_report)

        par11 = document.add_paragraph()
        run = par11.add_run(
            '\n    Значение СКО погрешности измерения угловых позиций Брэгговских отражений составляет ' + str(
                round(sko_int[2][1], 4)) + '.').font.bold = True

        document.save(filename.name)

        res = messagebox.askquestion('Отчёт сохранён', 'Открыть сгенерированный файл?')

        # if "yes", expands the search area and starts the process anew
        if res == 'yes':
            os.startfile(filename.name)
        elif res == 'no':
            return None
        else:
            messagebox.showwarning('error', 'Something went wrong!')

    # new interface buttons
    report_btn = Button(frame_for_down_buttons, text="Создать отчёт", width=20, command=make_report)
    another_file_btn = Button(frame_for_down_buttons, text="Выбрать другой файл", width=20, command=chose_another_file)
    info_btn = Button(frame_for_down_buttons, text="Информация", width=20, command=info_reload.reload)
    another_file_btn.pack(side=LEFT)
    info_btn.pack(side=LEFT)
    report_btn.pack(side=BOTTOM)

    # sets the size of the window and places it in the center of the screen
    secondary_window.update_idletasks()
    secondary_w = secondary_window.geometry()
    secondary_w = secondary_w.split('+')
    secondary_w = secondary_w[0].split('x')
    width_secondary_window = int(secondary_w[0])
    height_secondary_window = int(secondary_w[1])

    secondary_width = secondary_window.winfo_screenwidth()
    secondary_height = secondary_window.winfo_screenheight()
    secondary_width = secondary_width // 2
    secondary_height = secondary_height // 2
    secondary_width = secondary_width - width_secondary_window // 2
    secondary_height = secondary_height - height_secondary_window // 2
    secondary_window.geometry('+{}+{}'.format(secondary_width, secondary_height))


main_window = Tk()
main_window.title("Metrology")

# disables the ability to zoom the page
main_window.resizable(False, False)

# frame for the main interface
frame1 = LabelFrame(main_window)
frame1.pack(side=LEFT)

# outputs the information about the absolute error in the GUI
start_btn = Button(frame1, text="Выбрать \nи обработать файл  ", relief=GROOVE, width=18, command=open_file)
start_btn.pack(side=TOP)
search_area_button = Button(frame1, text="Область поиска\n данных в документе", relief=GROOVE, width=18,
                            command=change_search_area)
search_area_button.pack(side=TOP)

# sets the size of the window and places it in the center of the screen
main_window.update_idletasks()  # Updates information after all frames are created
s = main_window.geometry()
s = s.split('+')
s = s[0].split('x')
width_main_window = int(s[0])
height_main_window = int(s[1])

w = main_window.winfo_screenwidth()
h = main_window.winfo_screenheight()
w = w // 2
h = h // 2
w = w - width_main_window // 2
h = h - height_main_window // 2
main_window.geometry('+{}+{}'.format(w, h))

main_window.mainloop()
